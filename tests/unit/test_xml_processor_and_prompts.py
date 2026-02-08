from pathlib import Path

from docx import Document

from app.core.xml_processor import XMLProcessor
from app.prompts.rewrite_prompts import get_prompt


def test_xml_processor_generates_stable_paragraph_ids(tmp_path):
    doc_path = tmp_path / "stable_ids.docx"
    doc = Document()
    doc.add_paragraph("first")
    doc.add_paragraph("")
    doc.add_paragraph("second")
    doc.save(doc_path)

    processor = XMLProcessor(str(doc_path))
    skeleton = processor.extract_skeleton()
    ids = [item["id"] for item in skeleton["paragraphs"]]

    assert ids == ["para_000000", "para_000002"]


def test_xml_processor_apply_modifications_reports_failures(tmp_path):
    doc_path = tmp_path / "apply_modifications.docx"
    doc = Document()
    doc.add_paragraph("alpha")
    doc.add_paragraph("beta")
    doc.save(doc_path)

    processor = XMLProcessor(str(doc_path))
    processor.extract_skeleton()
    result = processor.apply_modifications(
        {
            "para_000000": "alpha_updated",
            "para_missing": "missing",
        }
    )

    assert result["applied_ids"] == []
    assert result["failed_ids"] == ["para_missing"]
    assert processor.paragraph_map["para_000000"].text == "alpha"


def test_prompt_supports_sentence_unit_and_option_count():
    prompt = get_prompt(
        mode="ai_detection",
        language="en",
        text="Original sentence.",
        unit="sentence",
        option_count=2,
    )

    assert "Rewrite only this sentence" in prompt
    assert "provide 2 different rewritten versions" in prompt
