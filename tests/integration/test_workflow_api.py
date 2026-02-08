from io import BytesIO

from docx import Document

from app.core.xml_processor import get_processor
import app.api.v1.rewrite as rewrite_api
import app.services.deepseek as deepseek_module


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def upload_sample_doc(client, sample_docx_bytes):
    response = client.post(
        "/api/v1/upload",
        files={"file": ("sample.docx", sample_docx_bytes, DOCX_MIME)},
    )
    assert response.status_code == 200
    return response.json()["data"]


def test_upload_ids_are_registered(client, sample_docx_bytes):
    uploaded = upload_sample_doc(client, sample_docx_bytes)
    processor = get_processor(uploaded["id"])

    paragraph_ids = [paragraph["id"] for paragraph in uploaded["paragraphs"]]
    assert paragraph_ids, "上传后至少应有一个非空段落"
    for paragraph_id in paragraph_ids:
        assert paragraph_id in processor.paragraph_map


def test_get_uploaded_document_success(client, sample_docx_bytes):
    uploaded = upload_sample_doc(client, sample_docx_bytes)

    response = client.get(f"/api/v1/upload/documents/{uploaded['id']}")

    assert response.status_code == 200
    payload = response.json()
    assert payload["success"] is True
    assert payload["data"]["id"] == uploaded["id"]
    assert payload["data"]["total_paragraphs"] == uploaded["total_paragraphs"]


def test_get_uploaded_document_not_found(client):
    response = client.get("/api/v1/upload/documents/doc_missing")

    assert response.status_code == 404
    assert "文档不存在" in response.json()["detail"]


def test_export_rejects_invalid_paragraph_id(client, sample_docx_bytes):
    uploaded = upload_sample_doc(client, sample_docx_bytes)

    response = client.post(
        "/api/v1/export",
        json={
            "doc_id": uploaded["id"],
            "modifications": {
                "para_not_exist": "replacement",
            },
        },
    )

    assert response.status_code == 400
    detail = response.json()["detail"]
    assert detail["failed_ids"] == ["para_not_exist"]
    assert detail["applied_ids"] == []


def test_export_stats_aggregates_failed_ids(client, sample_docx_bytes):
    uploaded = upload_sample_doc(client, sample_docx_bytes)
    second_doc = upload_sample_doc(client, sample_docx_bytes)

    client.post(
        "/api/v1/export",
        json={
            "doc_id": uploaded["id"],
            "modifications": {"para_not_exist": "replacement"},
        },
    )
    client.post(
        "/api/v1/export",
        json={
            "doc_id": second_doc["id"],
            "modifications": {"para_not_exist": "replacement", "para_also_missing": "replacement"},
        },
    )

    response = client.get("/api/v1/export/stats?window_minutes=60&top_n=10")

    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["failed_requests"] == 2
    assert payload["summary"]["failed_ids"] == 3
    assert len(payload["by_doc"]) == 2

    response_filtered = client.get(f"/api/v1/export/stats?window_minutes=60&doc_id={uploaded['id']}")
    assert response_filtered.status_code == 200
    filtered_payload = response_filtered.json()
    assert filtered_payload["summary"]["failed_requests"] == 1
    assert filtered_payload["summary"]["failed_ids"] == 1
    assert filtered_payload["by_doc"][0]["doc_id"] == uploaded["id"]


def test_export_stats_window_limit_boundaries(client):
    ok_response = client.get("/api/v1/export/stats?window_minutes=43200")
    assert ok_response.status_code == 200

    too_large_response = client.get("/api/v1/export/stats?window_minutes=43201")
    assert too_large_response.status_code == 422


def test_export_stats_storage_endpoint(client, sample_docx_bytes):
    uploaded = upload_sample_doc(client, sample_docx_bytes)
    client.post(
        "/api/v1/export",
        json={
            "doc_id": uploaded["id"],
            "modifications": {"para_not_exist": "replacement"},
        },
    )

    response = client.get("/api/v1/export/stats/storage")

    assert response.status_code == 200
    payload = response.json()
    assert payload["db_size_bytes"] > 0
    assert payload["event_count"] >= 1
    assert payload["oldest_event_at"] is not None
    assert payload["newest_event_at"] is not None
    assert payload["thresholds"]["warn_bytes"] > 0
    assert payload["thresholds"]["critical_bytes"] >= payload["thresholds"]["warn_bytes"]
    assert payload["level"] in {"ok", "warn", "critical"}


def test_export_applies_modifications_to_docx(client, sample_docx_bytes):
    uploaded = upload_sample_doc(client, sample_docx_bytes)
    target_id = uploaded["paragraphs"][0]["id"]
    marker = "E2E_MARKER_20260207"

    response = client.post(
        "/api/v1/export",
        json={
            "doc_id": uploaded["id"],
            "modifications": {
                target_id: marker,
            },
        },
    )

    assert response.status_code == 200
    exported_doc = Document(BytesIO(response.content))
    non_empty = [p.text for p in exported_doc.paragraphs if p.text.strip()]
    assert marker in non_empty


def test_export_mixed_ids_does_not_pollute_cached_document(client, sample_docx_bytes):
    uploaded = upload_sample_doc(client, sample_docx_bytes)
    target_id = uploaded["paragraphs"][0]["id"]
    original_text = uploaded["paragraphs"][0]["text"]
    mutated_text = "MUTATED_SHOULD_NOT_PERSIST"

    response = client.post(
        "/api/v1/export",
        json={
            "doc_id": uploaded["id"],
            "modifications": {
                target_id: mutated_text,
                "para_not_exist": "invalid",
            },
        },
    )
    assert response.status_code == 400
    detail = response.json()["detail"]
    assert detail["failed_ids"] == ["para_not_exist"]
    assert detail["applied_ids"] == []

    follow_up = client.get(f"/api/v1/export/{uploaded['id']}")
    assert follow_up.status_code == 200
    exported_doc = Document(BytesIO(follow_up.content))
    non_empty = [p.text for p in exported_doc.paragraphs if p.text.strip()]

    assert original_text in non_empty
    assert mutated_text not in non_empty


def test_rewrite_supports_sentence_unit_and_option_count(client, monkeypatch):
    class FakeRewriteService:
        def rewrite_text(self, text, mode, language, unit, option_count, max_retries=3):
            options = [f"{unit}-option-{idx}" for idx in range(1, option_count + 1)]
            sources = ["deepseek"] * option_count
            return {
                "options": options,
                "sources": sources,
                "diagnostics": {
                    "marian": {
                        "enabled": False,
                        "eligible": language == "en" and mode == "ai_detection",
                        "attempted": False,
                        "dependency_ready": False,
                        "used": False,
                        "status": "disabled",
                        "reason": "USE_MARIAN_MT=false",
                    }
                },
            }

    monkeypatch.setattr(rewrite_api, "get_deepseek_service", lambda: FakeRewriteService())

    response = client.post(
        "/api/v1/rewrite",
        json={
            "text": "This is one sentence.",
            "mode": "ai_detection",
            "language": "en",
            "unit": "sentence",
            "option_count": 2,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["unit"] == "sentence"
    assert len(payload["options"]) == 2
    assert payload["meta"][0]["source"] == "deepseek"
    assert payload["diagnostics"]["marian"]["status"] == "disabled"


def test_rewrite_marian_used_path(client, monkeypatch):
    service = deepseek_module.DeepSeekService.__new__(deepseek_module.DeepSeekService)
    monkeypatch.setattr(
        service,
        "_rewrite_with_deepseek",
        lambda text, mode, language, unit, option_count, max_retries: ["d1", "d2", "d3"],
    )
    monkeypatch.setattr(
        deepseek_module,
        "get_marian_runtime_info",
        lambda: {
            "enabled": True,
            "dependency_ready": True,
            "dependencies": {"transformers": True, "torch": True, "sentencepiece": True},
            "status": "enabled",
            "reason": None,
        },
    )

    class FakeMarianService:
        def back_translate_en(self, text):
            return "marian_noise_option"

    monkeypatch.setattr(deepseek_module, "get_marian_service", lambda: FakeMarianService())
    monkeypatch.setattr(rewrite_api, "get_deepseek_service", lambda: service)

    response = client.post(
        "/api/v1/rewrite",
        json={
            "text": "This is one sentence.",
            "mode": "ai_detection",
            "language": "en",
            "unit": "sentence",
            "option_count": 3,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["diagnostics"]["marian"]["status"] == "used"
    assert payload["diagnostics"]["marian"]["attempted"] is True
    assert payload["diagnostics"]["marian"]["used"] is True
    assert payload["meta"][-1]["source"] == "marian"


def test_rewrite_marian_generation_failed_path(client, monkeypatch):
    service = deepseek_module.DeepSeekService.__new__(deepseek_module.DeepSeekService)
    monkeypatch.setattr(
        service,
        "_rewrite_with_deepseek",
        lambda text, mode, language, unit, option_count, max_retries: ["d1", "d2", "d3"],
    )
    monkeypatch.setattr(
        deepseek_module,
        "get_marian_runtime_info",
        lambda: {
            "enabled": True,
            "dependency_ready": True,
            "dependencies": {"transformers": True, "torch": True, "sentencepiece": True},
            "status": "enabled",
            "reason": None,
        },
    )

    class BrokenMarianService:
        def back_translate_en(self, text):
            raise RuntimeError("mock marian failure")

    monkeypatch.setattr(deepseek_module, "get_marian_service", lambda: BrokenMarianService())
    monkeypatch.setattr(rewrite_api, "get_deepseek_service", lambda: service)

    response = client.post(
        "/api/v1/rewrite",
        json={
            "text": "This is one sentence.",
            "mode": "ai_detection",
            "language": "en",
            "unit": "sentence",
            "option_count": 3,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["diagnostics"]["marian"]["status"] == "generation_failed"
    assert payload["diagnostics"]["marian"]["attempted"] is True
    assert payload["diagnostics"]["marian"]["used"] is False
    assert payload["meta"] == [{"source": "deepseek"}, {"source": "deepseek"}, {"source": "deepseek"}]


def test_api_info_contains_marian_runtime(client):
    response = client.get("/api/v1/info")

    assert response.status_code == 200
    payload = response.json()
    assert "runtime" in payload
    assert "marian" in payload["runtime"]
    marian_runtime = payload["runtime"]["marian"]
    assert "dependency_ready" in marian_runtime
    assert "model_loaded" in marian_runtime
    assert "first_load_duration_ms" in marian_runtime
    assert "load_attempts" in marian_runtime
    assert "load_failures" in marian_runtime
    assert "generation_attempts" in marian_runtime
    assert "generation_failures" in marian_runtime
    assert "generation_failure_rate" in marian_runtime
    assert "last_generation_error" in marian_runtime
    assert isinstance(marian_runtime["generation_failure_rate"], float)
    assert 0.0 <= marian_runtime["generation_failure_rate"] <= 1.0
