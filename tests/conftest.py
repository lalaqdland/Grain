from io import BytesIO
from pathlib import Path
import sys

import pytest
from docx import Document
from fastapi.testclient import TestClient

ROOT_DIR = Path(__file__).resolve().parent.parent
BACKEND_DIR = ROOT_DIR / "backend"
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from main import app  # noqa: E402
from config import get_settings  # noqa: E402
import app.api.v1.upload as upload_api  # noqa: E402
import app.api.v1.export as export_api  # noqa: E402
from app.core.xml_processor import clear_processors  # noqa: E402
from app.core.document_registry import clear_documents  # noqa: E402
from app.core.export_observability import clear_export_failure_stats  # noqa: E402


@pytest.fixture
def client(tmp_path, monkeypatch):
    settings = get_settings()
    temp_dir = tmp_path / "temp_storage"
    temp_dir.mkdir(parents=True, exist_ok=True)

    monkeypatch.setattr(settings, "temp_storage_path", str(temp_dir))
    monkeypatch.setattr(upload_api, "settings", settings)
    monkeypatch.setattr(export_api, "settings", settings)

    clear_processors()
    clear_documents()
    clear_export_failure_stats()
    with TestClient(app) as test_client:
        yield test_client
    clear_processors()
    clear_documents()
    clear_export_failure_stats()


@pytest.fixture
def sample_docx_bytes():
    doc = Document()
    doc.add_paragraph("This is the first paragraph. It has two sentences.")
    doc.add_paragraph("第二段用于测试导出写回是否生效。")
    doc.add_paragraph("Third paragraph remains unchanged.")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
