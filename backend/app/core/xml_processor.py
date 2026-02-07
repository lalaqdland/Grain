"""
XML骨架处理器 - 用于保留Word文档格式
"""

from typing import Dict, Any
from docx import Document


class XMLProcessor:
    """XML骨架处理器"""
    
    def __init__(self, file_path: str):
        """
        初始化处理器
        
        Args:
            file_path: .docx文件路径
        """
        self.file_path = file_path
        self.document = Document(file_path)
        self.skeleton = {}  # 存储XML骨架
        self.paragraph_map = {}  # 段落ID到段落对象的映射
        
    def extract_skeleton(self) -> Dict[str, Any]:
        """
        提取XML骨架和段落索引
        
        Returns:
            包含骨架信息的字典
        """
        paragraphs_info = []
        self.paragraph_map = {}
        
        for idx, para in enumerate(self.document.paragraphs):
            # 跳过空段落
            if not para.text.strip():
                continue
            
            # 使用稳定ID，避免上传展示与导出写回映射错位
            para_id = self._build_paragraph_id(idx)
            
            # 保存段落对象引用
            self.paragraph_map[para_id] = para
            
            # 提取段落信息
            para_info = {
                "id": para_id,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "index": idx
            }
            
            paragraphs_info.append(para_info)
        
        self.skeleton = {
            "file_path": self.file_path,
            "paragraphs": paragraphs_info,
            "total_paragraphs": len(paragraphs_info)
        }
        
        return self.skeleton
    
    def replace_paragraph_text(self, para_id: str, new_text: str) -> bool:
        """
        替换段落文本（保留格式）
        
        Args:
            para_id: 段落ID
            new_text: 新文本
            
        Returns:
            是否成功
        """
        if para_id not in self.paragraph_map:
            return False
        
        para = self.paragraph_map[para_id]
        
        # 保留原有格式，只替换文本
        # 清空段落内容
        for run in para.runs:
            run.text = ""
        
        # 如果段落有runs，使用第一个run的格式
        if para.runs:
            para.runs[0].text = new_text
        else:
            # 如果没有runs，添加新的run
            para.add_run(new_text)
        
        return True
    
    @staticmethod
    def _build_paragraph_id(index: int) -> str:
        """根据段落索引生成稳定ID。"""
        return f"para_{index:06d}"
    
    def save_document(self, output_path: str) -> str:
        """
        保存修改后的文档
        
        Args:
            output_path: 输出路径
            
        Returns:
            保存的文件路径
        """
        self.document.save(output_path)
        return output_path
    
    def apply_modifications(self, modifications: Dict[str, str]) -> Dict[str, Any]:
        """
        批量应用修改
        
        Args:
            modifications: 段落ID到新文本的映射
        
        Returns:
            应用结果明细
        """
        requested_ids = list(modifications.keys())
        failed_ids = [para_id for para_id in requested_ids if para_id not in self.paragraph_map]
        if failed_ids:
            # 原子语义：存在无效ID时不执行任何修改。
            return {
                "requested_count": len(modifications),
                "applied_count": 0,
                "failed_count": len(failed_ids),
                "applied_ids": [],
                "failed_ids": failed_ids,
            }

        applied_ids = []
        original_text_map = {
            para_id: self.paragraph_map[para_id].text
            for para_id in requested_ids
        }
        try:
            for para_id, new_text in modifications.items():
                if not self.replace_paragraph_text(para_id, new_text):
                    raise RuntimeError(f"段落替换失败: {para_id}")
                applied_ids.append(para_id)
        except Exception:
            # 防御性回滚：确保处理器内存状态不会出现半应用结果。
            for para_id in applied_ids:
                self.replace_paragraph_text(para_id, original_text_map[para_id])
            raise

        return {
            "requested_count": len(modifications),
            "applied_count": len(applied_ids),
            "failed_count": 0,
            "applied_ids": applied_ids,
            "failed_ids": [],
        }


# 全局存储：文档ID到XMLProcessor的映射
_document_processors: Dict[str, XMLProcessor] = {}


def get_processor(doc_id: str) -> XMLProcessor:
    """
    获取文档处理器
    
    Args:
        doc_id: 文档ID
        
    Returns:
        XMLProcessor实例
        
    Raises:
        ValueError: 文档不存在
    """
    if doc_id not in _document_processors:
        raise ValueError(f"文档不存在: {doc_id}")
    return _document_processors[doc_id]


def register_processor(doc_id: str, processor: XMLProcessor) -> None:
    """
    注册文档处理器
    
    Args:
        doc_id: 文档ID
        processor: XMLProcessor实例
    """
    _document_processors[doc_id] = processor


def remove_processor(doc_id: str) -> None:
    """
    移除文档处理器
    
    Args:
        doc_id: 文档ID
    """
    if doc_id in _document_processors:
        del _document_processors[doc_id]


def clear_processors() -> None:
    """清理全部处理器（测试场景使用）。"""
    _document_processors.clear()

