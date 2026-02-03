"""
Word文档解析器
"""

import os
import uuid
from typing import List, Dict, Any
from docx import Document
from pathlib import Path


class DocxParser:
    """Word文档解析器"""
    
    def __init__(self, file_path: str):
        """
        初始化解析器
        
        Args:
            file_path: .docx文件路径
        """
        self.file_path = file_path
        self.document = None
        
    def parse(self) -> Dict[str, Any]:
        """
        解析Word文档
        
        Returns:
            包含文档信息的字典
            
        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件格式错误
        """
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"文件不存在: {self.file_path}")
        
        try:
            self.document = Document(self.file_path)
        except Exception as e:
            raise ValueError(f"无法解析文档: {str(e)}")
        
        # 提取段落
        paragraphs = self._extract_paragraphs()
        
        # 获取文件名
        filename = Path(self.file_path).name
        
        return {
            "filename": filename,
            "paragraphs": paragraphs,
            "total_paragraphs": len(paragraphs)
        }
    
    def _extract_paragraphs(self) -> List[Dict[str, Any]]:
        """
        提取文档段落
        
        Returns:
            段落列表
        """
        paragraphs = []
        
        for para in self.document.paragraphs:
            # 跳过空段落
            if not para.text.strip():
                continue
            
            # 生成唯一ID
            para_id = self._generate_id()
            
            # 提取段落信息
            para_info = {
                "id": para_id,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "is_modified": False,
                "original_text": None
            }
            
            paragraphs.append(para_info)
        
        return paragraphs
    
    @staticmethod
    def _generate_id() -> str:
        """
        生成唯一ID
        
        Returns:
            UUID字符串
        """
        return f"para_{uuid.uuid4().hex[:12]}"
    
    def get_paragraph_count(self) -> int:
        """
        获取段落数量
        
        Returns:
            段落数量
        """
        if not self.document:
            return 0
        return len([p for p in self.document.paragraphs if p.text.strip()])

